from pathlib import Path

import pytest

pytest.importorskip("pypdf")
pytest.importorskip("docx")

from docx import Document as DocxDocument

from contextos.ingestion.documents import DocumentExtractor


def _write_minimal_pdf(path: Path, text: str) -> None:
    """Builds a real, minimal-but-valid single-page PDF by hand (no PDF-writing
    library needed) so the test exercises pypdf's actual parser, not a mock."""
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
            b"/MediaBox [0 0 300 300] /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    stream = f"BT /F1 18 Tf 20 150 Td ({text}) Tj ET".encode()
    objects.append(b"<< /Length %d >>\nstream\n" % len(stream) + stream + b"\nendstream")

    buf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(buf))
        buf += f"{index} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_offset = len(buf)
    buf += f"xref\n0 {len(objects) + 1}\n".encode()
    buf += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        buf += f"{offset:010d} 00000 n \n".encode()
    buf += f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode()
    path.write_bytes(bytes(buf))


@pytest.mark.asyncio
async def test_extracts_text_from_a_real_pdf(tmp_path: Path) -> None:
    pdf_path = tmp_path / "report.pdf"
    _write_minimal_pdf(pdf_path, "Hello from a real PDF")

    nodes = await DocumentExtractor(pdf_path).extract(tenant_id="t1")

    assert len(nodes) == 1
    assert "Hello from a real PDF" in (nodes[0].content or "")
    assert nodes[0].metadata["source_type"] == "document"
    assert nodes[0].metadata["source_path"] == str(pdf_path)


@pytest.mark.asyncio
async def test_extracts_text_from_a_real_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "notes.docx"
    document = DocxDocument()
    document.add_paragraph("First paragraph from a real docx file.")
    document.add_paragraph("Second paragraph, different sentence.")
    document.save(str(docx_path))

    nodes = await DocumentExtractor(docx_path, chunk_chars=10_000).extract(tenant_id="t1")

    assert len(nodes) == 1
    assert "First paragraph from a real docx file." in (nodes[0].content or "")
    assert "Second paragraph, different sentence." in (nodes[0].content or "")


@pytest.mark.asyncio
async def test_plain_text_file_round_trips_unchanged(tmp_path: Path) -> None:
    text_path = tmp_path / "readme.txt"
    text_path.write_text("Just a plain text file.\n\nWith two paragraphs.", encoding="utf-8")

    nodes = await DocumentExtractor(text_path).extract(tenant_id="t1")

    assert len(nodes) == 1
    assert nodes[0].content == "Just a plain text file.\n\nWith two paragraphs."
    assert nodes[0].title == "readme.txt"


@pytest.mark.asyncio
async def test_long_document_is_split_into_multiple_chunks(tmp_path: Path) -> None:
    paragraphs = [f"Paragraph number {i} with some real sentence content in it." for i in range(20)]
    text_path = tmp_path / "long.md"
    text_path.write_text("\n\n".join(paragraphs), encoding="utf-8")

    nodes = await DocumentExtractor(text_path, chunk_chars=200).extract(tenant_id="t1")

    assert len(nodes) > 1
    assert all(node.metadata["chunk_count"] == len(nodes) for node in nodes)
    assert [node.metadata["chunk_index"] for node in nodes] == list(range(len(nodes)))
    assert "part 1" in (nodes[0].title or "")
    # Every original paragraph shows up in exactly the reassembled chunks, nothing lost.
    reassembled = "\n\n".join(node.content or "" for node in nodes)
    for paragraph in paragraphs:
        assert paragraph in reassembled


@pytest.mark.asyncio
async def test_empty_file_yields_no_nodes(tmp_path: Path) -> None:
    text_path = tmp_path / "empty.txt"
    text_path.write_text("", encoding="utf-8")

    nodes = await DocumentExtractor(text_path).extract(tenant_id="t1")

    assert nodes == []


def test_unsupported_extension_raises() -> None:
    with pytest.raises(ValueError, match="Unsupported document type"):
        DocumentExtractor("archive.zip")


@pytest.mark.asyncio
async def test_ingest_source_persists_every_extracted_node(tmp_path: Path) -> None:
    from contextos import ContextOS

    text_path = tmp_path / "notes.txt"
    text_path.write_text("A single short note.", encoding="utf-8")

    context_os = ContextOS()
    nodes = await context_os.ingest_source(DocumentExtractor(text_path), tenant_id="t1")

    assert len(nodes) == 1
    stored = await context_os.store.get_node("t1", nodes[0].id)
    assert stored is not None
    assert stored.content == "A single short note."
